"""
Ekstrak properti formatting dari file DOCX menggunakan python-docx.

Digunakan oleh: validator.py, rule_validator.py

Tujuan: Membaca semua properti formatting (typography, page layout, spacing,
document structure) dari file DOCX untuk dibandingkan dengan rules.
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from model_ai.validation.models import DocxProperties


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Blok konstanta `PAPER_SIZE_MAP` untuk menyimpan konfigurasi/registry yang dipakai berulang.
# ---------------------------------------------------------------------------
PAPER_SIZE_MAP: dict[tuple[float, float], str] = {
    (21.0, 29.7): "A4",
    (21.0, 33.0): "F4",
    (14.85, 21.0): "A5",
    (29.7, 42.0): "A3",
    (21.59, 27.94): "LETTER",
}


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_get_paper_size` untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _get_paper_size(width_cm: float, height_cm: float) -> str:
    """Map paper dimensions to standard paper size names."""
    # Check portrait (width < height)
    if width_cm < height_cm:
        key = (round(width_cm, 1), round(height_cm, 1))
        if key in PAPER_SIZE_MAP:
            return PAPER_SIZE_MAP[key]
    # Check landscape (width > height)
    else:
        key = (round(height_cm, 1), round(width_cm, 1))
        if key in PAPER_SIZE_MAP:
            return PAPER_SIZE_MAP[key]
    # Return approximate size if not in map
    return f"{round(width_cm, 1)}x{round(height_cm, 1)}cm"


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_get_alignment_string` untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _get_alignment_string(alignment) -> str:
    """Convert python-docx alignment enum to string."""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    }
    return mapping.get(alignment, "LEFT")


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_get_orientation_string" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _get_orientation_string(orientation) -> str:
    """Convert python-docx orientation enum to string."""
    mapping = {
        WD_ORIENT.PORTRAIT: "PORTRAIT",
        WD_ORIENT.LANDSCAPE: "LANDSCAPE",
    }
    return mapping.get(orientation, "PORTRAIT")


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_get_line_spacing_rule" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _get_line_spacing_rule(line_spacing) -> str:
    """Determine line spacing rule from the value."""
    if line_spacing is None:
        return "UNKNOWN"
    # python-docx returns LineSpacing value in twips (1/20 of a point)
    # or a float for MULTIPLE, or None for AT_LEAST
    return "MULTIPLE"  # Most common for documents


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_extract_typography" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _extract_typography(doc: Document) -> dict:
    """Extract typography properties from document."""
    # Get default font from Normal style
    normal_style = doc.styles["Normal"]
    font_family = normal_style.font.name
    if font_family is None:
        font_family = "Unknown"

    font_size_pt = None
    if normal_style.font.size:
        font_size_pt = int(normal_style.font.size.pt)

    # Check heading styles
    heading_bold = None
    heading_all_caps = None
    heading_font_size = None

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[style_name]
            if heading_font_size is None and style.font.size:
                heading_font_size = int(style.font.size.pt)
            if heading_bold is None and style.font.bold:
                heading_bold = style.font.bold
            if heading_all_caps is None and style.font.all_caps:
                heading_all_caps = style.font.all_caps
        except KeyError:
            continue

    # Sample runs from document to get average font info
    run_count = 0
    fonts_found: set[str] = set()
    sizes_found: set[int] = set()

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                run_count += 1
                if run.font.name and run.font.name != font_family:
                    fonts_found.add(run.font.name)
                try:
                    if run.font.size:
                        sizes_found.add(int(run.font.size.pt))
                except (ValueError, TypeError):
                    pass

    return {
        "font_family": font_family,
        "font_size_body_pt": font_size_pt,
        "font_size_heading_pt": heading_font_size,
        "heading_bold": heading_bold,
        "heading_all_caps": heading_all_caps,
    }


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_extract_page_layout" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _twips_to_cm(value_str: str | None) -> float | None:
    """Convert twips string (may be float) to cm. 1 inch = 1440 twips = 2.54 cm."""
    if value_str is None:
        return None
    try:
        return float(value_str) / 1440 * 2.54
    except (ValueError, TypeError):
        return None


def _extract_page_layout(doc: Document) -> dict:
    """Extract page layout properties from first section."""
    section = doc.sections[0]

    # Read margins directly from XML to avoid python-docx int() cast on float strings
    pgMar = section._sectPr.find(qn("w:pgMar"))
    if pgMar is not None:
        margin_top    = _twips_to_cm(pgMar.get(qn("w:top")))
        margin_bottom = _twips_to_cm(pgMar.get(qn("w:bottom")))
        margin_left   = _twips_to_cm(pgMar.get(qn("w:left")))
        margin_right  = _twips_to_cm(pgMar.get(qn("w:right")))
    else:
        margin_top = margin_bottom = margin_left = margin_right = None

    # Paper size — read from XML for same reason
    pgSz = section._sectPr.find(qn("w:pgSz"))
    if pgSz is not None:
        page_width_cm  = _twips_to_cm(pgSz.get(qn("w:w")))
        page_height_cm = _twips_to_cm(pgSz.get(qn("w:h")))
    else:
        page_width_cm = page_height_cm = None

    paper_size = None
    if page_width_cm and page_height_cm:
        paper_size = _get_paper_size(page_width_cm, page_height_cm)

    # Orientation — read from XML to avoid python-docx enum parse issues
    orient_attr = pgSz.get(qn("w:orient")) if pgSz is not None else None
    if orient_attr == "landscape":
        orientation = "landscape"
    else:
        orientation = "portrait"

    return {
        "margin_top_cm": round(margin_top, 2) if margin_top else None,
        "margin_bottom_cm": round(margin_bottom, 2) if margin_bottom else None,
        "margin_left_cm": round(margin_left, 2) if margin_left else None,
        "margin_right_cm": round(margin_right, 2) if margin_right else None,
        "paper_size": paper_size,
        "orientation": orientation,
    }


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_extract_spacing" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _extract_spacing(doc: Document) -> dict:
    """Extract spacing properties from document."""
    # Get from Normal style
    normal_style = doc.styles["Normal"]

    line_spacing = None
    if normal_style.paragraph_format.line_spacing:
        # Line spacing can be in points (float) or LINE_SPACING enum
        ls_value = normal_style.paragraph_format.line_spacing
        if isinstance(ls_value, (int, float)):
            line_spacing = round(float(ls_value), 2)
        else:
            line_spacing = 1.15  # default

    paragraph_alignment = _get_alignment_string(normal_style.paragraph_format.alignment)

    first_line_indent = None
    if normal_style.paragraph_format.first_line_indent:
        first_line_indent = normal_style.paragraph_format.first_line_indent.cm

    # Check references style for hanging indent
    references_hanging_indent = None
    try:
        ref_style = doc.styles["Citations"] or doc.styles["Bibliography"]
        if ref_style and ref_style.paragraph_format:
            # Check if left indent > first line indent (indicates hanging indent)
            left_indent = ref_style.paragraph_format.left_indent
            if left_indent and first_line_indent is not None:
                references_hanging_indent = left_indent.cm > first_line_indent
    except (KeyError, AttributeError):
        pass

    return {
        "line_spacing": line_spacing,
        "line_spacing_rule": "MULTIPLE",
        "paragraph_alignment": paragraph_alignment,
        "first_line_indent_cm": round(first_line_indent, 2) if first_line_indent else None,
        "references_hanging_indent": references_hanging_indent,
    }


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_extract_document_structure" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _extract_document_structure(doc: Document) -> dict:
    """Extract document structure (headings, sections)."""
    # Count headings by style
    heading_count = 0
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            heading_count += 1

    # Count sections (physical sections in document)
    section_count = len(doc.sections)

    # Check first few paragraphs for preliminary pages
    has_cover = False
    has_approval = False
    has_summary = False

    first_paras = [p.text.strip().upper() for p in doc.paragraphs[:20] if p.text.strip()]

    for text in first_paras:
        if "SAMPUL" in text or "COVER" in text:
            has_cover = True
        if "PENGESAHAN" in text or "PERSETUJUAN" in text or "APPROVAL" in text:
            has_approval = True
        if "RINGKASAN" in text or "ABSTRACT" in text:
            has_summary = True

    return {
        "heading_count": heading_count,
        "section_count": section_count,
        "has_halaman_sampul": has_cover,
        "has_halaman_pengesahan": has_approval,
        "has_ringkasan": has_summary,
    }


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_extract_figures_tables" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _extract_figures_tables(doc: Document) -> dict:
    """Extract figures and tables information."""
    table_count = len(doc.tables)

    # Estimate figure count by checking image elements
    figure_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            figure_count += 1

    return {
        "table_count": table_count,
        "figure_count": figure_count,
    }


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `_extract_numbering" untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def _extract_numbering(doc: Document) -> dict:
    """Extract numbering properties."""
    # Page numbering is typically in headers/footers
    # For now, return None as this requires more complex footer parsing

    chapter_format = "BAB {n}"  # Default assumption for PKM proposal

    return {
        "chapter_format": chapter_format,
        "preliminary_page_format": None,  # Requires footer parsing
        "content_page_format": None,  # Requires footer parsing
    }


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `extract_docx_properties` untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def extract_docx_properties(docx_path: str | Path) -> DocxProperties:
    """Extract all formatting properties from a DOCX file.

    Args:
        docx_path: Path to the DOCX file to extract properties from.

    Returns:
        DocxProperties object containing all extracted properties.

    Raises:
        FileNotFoundError: If the DOCX file does not exist.
        Exception: If there is an error reading the DOCX file.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {e}")

    # Extract all properties
    typography = _extract_typography(doc)
    page_layout = _extract_page_layout(doc)
    spacing = _extract_spacing(doc)
    doc_structure = _extract_document_structure(doc)
    figures_tables = _extract_figures_tables(doc)
    numbering = _extract_numbering(doc)

    # Combine all into DocxProperties
    props = DocxProperties(
        # Typography
        font_family=typography["font_family"],
        font_size_body_pt=typography["font_size_body_pt"],
        font_size_heading_pt=typography["font_size_heading_pt"],
        heading_bold=typography["heading_bold"],
        heading_all_caps=typography["heading_all_caps"],
        # Page Layout
        margin_top_cm=page_layout["margin_top_cm"],
        margin_bottom_cm=page_layout["margin_bottom_cm"],
        margin_left_cm=page_layout["margin_left_cm"],
        margin_right_cm=page_layout["margin_right_cm"],
        paper_size=page_layout["paper_size"],
        orientation=page_layout["orientation"],
        # Spacing
        line_spacing=spacing["line_spacing"],
        line_spacing_rule=spacing["line_spacing_rule"],
        paragraph_alignment=spacing["paragraph_alignment"],
        first_line_indent_cm=spacing["first_line_indent_cm"],
        references_hanging_indent=spacing["references_hanging_indent"],
        # Document Structure
        heading_count=doc_structure["heading_count"],
        section_count=doc_structure["section_count"],
        has_halaman_sampul=doc_structure["has_halaman_sampul"],
        has_halaman_pengesahan=doc_structure["has_halaman_pengesahan"],
        has_ringkasan=doc_structure["has_ringkasan"],
        # Numbering
        chapter_format=numbering["chapter_format"],
        preliminary_page_format=numbering["preliminary_page_format"],
        content_page_format=numbering["content_page_format"],
        # Figures & Tables
        table_count=figures_tables["table_count"],
        figure_count=figures_tables["figure_count"],
    )

    return props


# ---------------------------------------------------------------------------
# Digunakan oleh: Dipakai internal di file ini atau dipanggil dari entrypoint runtime.
# Mendefinisikan fungsi `extract_docx_properties_dict` untuk kebutuhan modul `docx_property_extractor`.
# ---------------------------------------------------------------------------
def extract_docx_properties_dict(docx_path: str | Path) -> dict:
    """Extract all formatting properties from a DOCX file as dict.

    Args:
        docx_path: Path to the DOCX file to extract properties from.

    Returns:
        Dictionary containing all extracted properties.
    """
    props = extract_docx_properties(docx_path)
    return props.to_dict()