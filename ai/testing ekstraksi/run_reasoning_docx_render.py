from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
APP_DIR = SCRIPT_DIR.parent
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))
DATA_DIR = APP_DIR / "data"
DEFAULT_CHUNKS = DATA_DIR / "output_chunks.json"
DEFAULT_REASONING_REPORT = SCRIPT_DIR / "reasoning_translation_report.json"
DEFAULT_DOCX_OUTPUT = SCRIPT_DIR / "reasoning_rendered_proposal.docx"
DEFAULT_RENDER_REPORT = SCRIPT_DIR / "reasoning_docx_render_report.json"
BLACK = RGBColor(0, 0, 0)


def _load_json(path: Path) -> Any:
    if not path.exists():
        raise FileNotFoundError(f"File tidak ditemukan: {path}")
    return json.loads(path.read_text(encoding="utf-8-sig"))


def _load_metadata_payload(source_doc: str) -> dict[str, Any]:
    from model_ai.metadata_repository import load_document_metadata_payload

    return load_document_metadata_payload(source_doc)


def _normalize_text(value: Any) -> str:
    text = str(value or "").lower()
    text = re.sub(r"[*_`#|:().,\-/]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _tokens(value: Any) -> set[str]:
    ignored = {"bab", "daftar", "dan", "yang", "jika", "ada", "kegiatan"}
    return {token for token in _normalize_text(value).split() if token and token not in ignored}


def _accepted_results(report: dict[str, Any]) -> dict[str, dict[str, Any]]:
    results = report.get("results", [])
    if not isinstance(results, list):
        return {}
    return {
        str(item.get("source_field")): item
        for item in results
        if isinstance(item, dict) and item.get("final_status") == "accepted"
    }


def _value(mapping: dict[str, dict[str, Any]], source_field: str, default: Any = None) -> Any:
    item = mapping.get(source_field)
    if not item:
        return default
    value = item.get("normalized_value")
    return default if value is None else value


def _set_font_name(font: Any, font_name: str) -> None:
    font.name = font_name
    if font._element.rPr is not None:
        font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_black_font(font: Any) -> None:
    font.color.rgb = BLACK


def _set_black_run(run: Any) -> None:
    run.font.color.rgb = BLACK


def _alignment(value: Any) -> Any:
    normalized = str(value or "").strip().upper()
    if normalized == "JUSTIFY":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if normalized == "CENTER":
        return WD_ALIGN_PARAGRAPH.CENTER
    if normalized == "RIGHT":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _page_label(chunk: dict[str, Any]) -> str:
    page = chunk.get("page") if isinstance(chunk.get("page"), dict) else {}
    start = page.get("start")
    end = page.get("end")
    if start and end and start != end:
        return f"Hal. {start}-{end}"
    if start:
        return f"Hal. {start}"
    return "Hal. tidak tersedia"


def _source_line(chunk: dict[str, Any]) -> str:
    header = str(chunk.get("chunk_parent") or "Tanpa header").strip()
    return f"Sumber: {_page_label(chunk)} | Header: {header}"


def _clean_chunk_snippet(text: str, *, limit: int = 520) -> str:
    cleaned = re.sub(r"<br\s*/?>", " ", text, flags=re.IGNORECASE)
    cleaned = re.sub(r"\|[-:| ]+\|", " ", cleaned)
    cleaned = re.sub(r"[*_`#]+", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rsplit(" ", 1)[0] + "..."


def _sentence_case(text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        return cleaned
    return cleaned[0].upper() + cleaned[1:]


def _rewrite_chunk_as_instruction(section_spec: dict[str, Any], chunk: dict[str, Any]) -> str:
    section_type = str(section_spec.get("type", ""))
    title = str(section_spec.get("title") or section_type.replace("_", " "))
    snippet = _clean_chunk_snippet(str(chunk.get("content", "")), limit=360)

    if section_type == "bab":
        return (
            f"Gunakan informasi sumber ini untuk menyusun bagian {title}. "
            f"Uraikan substansi utama dengan bahasa proposal yang runtut, bukan menyalin mentah. "
            f"Pokok yang perlu dibawa ke naskah: {_sentence_case(snippet)}"
        )
    if section_type == "daftar_pustaka":
        return (
            "Gunakan aturan ini untuk menyusun daftar pustaka. Pastikan setiap referensi yang muncul "
            f"di isi proposal juga tercantum di daftar pustaka. Ringkasan aturan sumber: {_sentence_case(snippet)}"
        )
    if section_type.startswith("daftar_"):
        return (
            f"Gunakan sumber ini untuk menentukan apakah {title} perlu diisi. "
            f"Jika elemen terkait tidak ada, bagian ini dapat dibiarkan sebagai daftar kosong yang rapi. "
            f"Dasar instruksi: {_sentence_case(snippet)}"
        )
    return f"Gunakan sumber ini sebagai arahan penulisan bagian {title}: {_sentence_case(snippet)}"


def _section_query_text(section_spec: dict[str, Any]) -> str:
    section_type = str(section_spec.get("type", ""))
    title = str(section_spec.get("title", ""))
    if section_type == "bab":
        return f"BAB {section_spec.get('number')} {title}"
    return f"{section_type} {title}"


def _is_proposal_structure_chunk(chunk: dict[str, Any]) -> bool:
    header_norm = _normalize_text(chunk.get("chunk_parent", ""))
    if "laporan kemajuan" in header_norm or "laporan akhir" in header_norm:
        return False

    page = chunk.get("page") if isinstance(chunk.get("page"), dict) else {}
    page_start = page.get("start")
    if isinstance(page_start, int) and 8 <= page_start <= 11:
        return True
    return "sistematika penulisan proposal" in header_norm


def _score_chunk_for_section(section_spec: dict[str, Any], chunk: dict[str, Any]) -> float:
    if not _is_proposal_structure_chunk(chunk):
        return 0.0

    query = _section_query_text(section_spec)
    query_norm = _normalize_text(query)
    chunk_header = str(chunk.get("chunk_parent", ""))
    chunk_content = str(chunk.get("content", ""))
    chunk_norm = _normalize_text(f"{chunk_header} {chunk_content}")
    section_type = str(section_spec.get("type", ""))
    title_norm = _normalize_text(section_spec.get("title", ""))

    if section_type.startswith("daftar_") and title_norm and title_norm not in chunk_norm:
        return 0.0

    score = 0.0
    query_tokens = _tokens(query)
    chunk_tokens = _tokens(chunk_norm)
    if query_tokens:
        score += len(query_tokens & chunk_tokens) / len(query_tokens)
        if "biaya" in query_tokens and "biaya" in chunk_tokens:
            score += 0.35
        if "jadwal" in query_tokens and "jadwal" in chunk_tokens:
            score += 0.35
        header_norm = _normalize_text(chunk_header)
        if "biaya" in query_tokens and "anggaran biaya" in header_norm:
            score += 0.85
        if "jadwal" in query_tokens and "jadwal" in header_norm:
            score += 1.2

    if section_type == "bab":
        bab_phrase = _normalize_text(f"bab {section_spec.get('number')}")
        if bab_phrase and bab_phrase in chunk_norm:
            score += 2.0
    if query_norm and query_norm in chunk_norm:
        score += 2.0

    title = str(section_spec.get("title", ""))
    if title_norm and title_norm in _normalize_text(chunk_header):
        score += 1.5

    proposal_hint = "sistematika penulisan proposal"
    if proposal_hint in _normalize_text(chunk_header):
        score += 0.4
    if "laporan kemajuan" in _normalize_text(chunk_header) or "laporan akhir" in _normalize_text(chunk_header):
        score -= 1.0

    return score


def _match_chunks_for_section(
    section_spec: dict[str, Any],
    chunks: list[dict[str, Any]],
    *,
    limit: int = 3,
) -> list[dict[str, Any]]:
    scored = [
        (_score_chunk_for_section(section_spec, chunk), chunk)
        for chunk in chunks
        if isinstance(chunk, dict) and str(chunk.get("content", "")).strip()
    ]
    scored.sort(key=lambda item: item[0], reverse=True)

    matched: list[dict[str, Any]] = []
    seen_sources: set[tuple[Any, Any, str]] = set()
    for score, chunk in scored:
        if score <= 0.35:
            continue
        page = chunk.get("page") if isinstance(chunk.get("page"), dict) else {}
        source_key = (page.get("start"), page.get("end"), str(chunk.get("chunk_parent") or ""))
        if source_key in seen_sources:
            continue
        seen_sources.add(source_key)
        matched.append(chunk)
        if len(matched) >= limit:
            break
    return matched


def _apply_page_settings(document: Document, mapping: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    applied: list[dict[str, Any]] = []
    section = document.sections[0]
    margin_fields = {
        "page_layout.margin_top_cm": ("top_margin", "section.top_margin"),
        "page_layout.margin_bottom_cm": ("bottom_margin", "section.bottom_margin"),
        "page_layout.margin_left_cm": ("left_margin", "section.left_margin"),
        "page_layout.margin_right_cm": ("right_margin", "section.right_margin"),
    }

    for source_field, (attribute, target_path) in margin_fields.items():
        margin_value = _value(mapping, source_field)
        if isinstance(margin_value, (int, float)):
            setattr(section, attribute, Cm(float(margin_value)))
            applied.append(
                {
                    "source_field": source_field,
                    "target_path": target_path,
                    "applied_value": margin_value,
                    "python_docx_expression": f"Cm({margin_value})",
                }
            )

    orientation = str(_value(mapping, "page_layout.orientation", "PORTRAIT")).upper()
    if orientation == "LANDSCAPE":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        applied.append(
            {
                "source_field": "page_layout.orientation",
                "target_path": "section.orientation",
                "applied_value": "LANDSCAPE",
                "python_docx_expression": "WD_ORIENT.LANDSCAPE",
            }
        )

    return applied


def _apply_styles(document: Document, mapping: dict[str, dict[str, Any]]) -> dict[str, Any]:
    font_name = str(_value(mapping, "typography.font_family", "Times New Roman"))
    body_size = float(_value(mapping, "typography.font_size_body_pt", 12))
    heading_size = float(_value(mapping, "typography.font_size_heading_pt", body_size))
    heading_bold = bool(_value(mapping, "typography.heading_bold", True))
    heading_all_caps = bool(_value(mapping, "typography.heading_all_caps", True))
    line_spacing = _value(mapping, "spacing.line_spacing", 1.15)
    paragraph_alignment = _alignment(_value(mapping, "spacing.paragraph_alignment", "JUSTIFY"))

    normal = document.styles["Normal"]
    _set_font_name(normal.font, font_name)
    _set_black_font(normal.font)
    normal.font.size = Pt(body_size)
    normal.paragraph_format.line_spacing = float(line_spacing)
    normal.paragraph_format.alignment = paragraph_alignment

    heading = document.styles["Heading 1"]
    _set_font_name(heading.font, font_name)
    _set_black_font(heading.font)
    heading.font.size = Pt(heading_size)
    heading.font.bold = heading_bold
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return {
        "font_name": font_name,
        "body_size_pt": body_size,
        "heading_size_pt": heading_size,
        "heading_bold": heading_bold,
        "heading_all_caps": heading_all_caps,
        "line_spacing": line_spacing,
        "paragraph_alignment": str(_value(mapping, "spacing.paragraph_alignment", "JUSTIFY")),
    }


def _heading_text(section_spec: dict[str, Any], *, all_caps: bool) -> str:
    section_type = str(section_spec.get("type", "section"))
    title = str(section_spec.get("title") or section_type.replace("_", " ").upper())
    if section_type == "bab":
        number = section_spec.get("number")
        title = f"BAB {number}. {title}" if number is not None else f"BAB. {title}"
    return title.upper() if all_caps else title


def _bookmark_name(index: int, section_spec: dict[str, Any]) -> str:
    section_type = re.sub(r"[^A-Za-z0-9_]+", "_", str(section_spec.get("type", "section")))
    number = section_spec.get("number")
    suffix = f"_{number}" if number is not None else ""
    return f"section_{index + 1}_{section_type}{suffix}"


def _add_bookmark(paragraph: Any, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_hyperlink(paragraph: Any, *, anchor: str, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(color)
    run_properties.append(underline)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_properties)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_center_heading(document: Document, text: str, *, bookmark_name: str | None = None, bookmark_id: int = 0) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    _set_black_run(run)
    if bookmark_name:
        _add_bookmark(paragraph, bookmark_name, bookmark_id)


def _add_body_paragraph(document: Document, text: str, *, alignment: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    _set_black_run(run)


def _section_logical_page(index: int) -> int:
    return index + 1


def _content_width_cm(document: Document) -> float:
    section = document.sections[0]
    width = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    return max(width, 10.0)


def _add_toc_entry(document: Document, *, title: str, anchor: str, page_number: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(_content_width_cm(document)),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.DOTS,
    )
    _add_internal_hyperlink(paragraph, anchor=anchor, text=title)
    visible_dots = "." * max(8, 62 - len(title) - len(str(page_number)))
    dots_run = paragraph.add_run(f" {visible_dots} ")
    _set_black_run(dots_run)
    paragraph.add_run("\t")
    _add_internal_hyperlink(paragraph, anchor=anchor, text=str(page_number))


def _render_daftar_isi(
    document: Document,
    sections: list[dict[str, Any]],
    *,
    anchors: dict[int, str],
    all_caps: bool,
) -> None:
    note = document.add_paragraph()
    note_run = note.add_run("Klik judul atau nomor halaman untuk berpindah ke bagian terkait.")
    note_run.italic = True
    note_run.font.size = Pt(10)
    _set_black_run(note_run)

    for index, item in enumerate(sections):
        section_type = str(item.get("type", ""))
        if section_type == "daftar_isi":
            continue
        title = _heading_text(item, all_caps=all_caps)
        _add_toc_entry(
            document,
            title=title,
            anchor=anchors[index],
            page_number=_section_logical_page(index),
        )


def _add_instructional_placeholder(
    document: Document,
    *,
    section_spec: dict[str, Any],
    matched_chunks: list[dict[str, Any]],
    alignment: Any,
) -> None:
    title = str(section_spec.get("title", section_spec.get("type", "bagian")))
    intro = (
        f"Instructional placeholder untuk {title}. Setiap blok di bawah menunjukkan chunk "
        "yang disiapkan sebagai bahan konteks LLM, lalu rewrite instruksi yang lebih siap "
        "dipakai untuk menulis bagian dokumen."
    )
    _add_body_paragraph(document, intro, alignment=alignment)

    if not matched_chunks:
        _add_body_paragraph(
            document,
            "Belum ditemukan chunk yang cukup relevan dari output_chunks.json untuk bagian ini.",
            alignment=alignment,
        )
        return

    for index, chunk in enumerate(matched_chunks, start=1):
        snippet = _clean_chunk_snippet(str(chunk.get("content", "")))

        context_title = document.add_paragraph()
        context_title.alignment = alignment
        context_number = context_title.add_run(f"{index}. Chunk yang dioper ke LLM")
        context_number.bold = True
        _set_black_run(context_number)

        context = document.add_paragraph()
        context.alignment = alignment
        context_run = context.add_run(f"Informasi chunk: {snippet}")
        _set_black_run(context_run)

        rewritten = document.add_paragraph()
        rewritten.alignment = alignment
        rewritten_label = rewritten.add_run("Rewrite instructional placeholder: ")
        rewritten_label.bold = True
        _set_black_run(rewritten_label)
        rewritten_run = rewritten.add_run(_rewrite_chunk_as_instruction(section_spec, chunk))
        _set_black_run(rewritten_run)

        source = document.add_paragraph()
        source.alignment = alignment
        source_run = source.add_run(_source_line(chunk))
        source_run.italic = True
        source_run.font.size = Pt(10)
        _set_black_run(source_run)


def _render_section_body(
    document: Document,
    section_spec: dict[str, Any],
    *,
    sections: list[dict[str, Any]],
    anchors: dict[int, str],
    all_caps: bool,
    matched_chunks: list[dict[str, Any]],
    alignment: Any,
) -> None:
    section_type = str(section_spec.get("type", "section"))
    title = str(section_spec.get("title", section_type))

    if section_type == "daftar_isi":
        _render_daftar_isi(document, sections, anchors=anchors, all_caps=all_caps)
        return
    if section_type in {"daftar_gambar", "daftar_tabel", "daftar_lampiran"}:
        _add_instructional_placeholder(
            document,
            section_spec=section_spec,
            matched_chunks=matched_chunks,
            alignment=alignment,
        )
        return
    if section_type == "daftar_pustaka":
        _add_instructional_placeholder(
            document,
            section_spec=section_spec,
            matched_chunks=matched_chunks,
            alignment=alignment,
        )
        return

    _add_instructional_placeholder(
        document,
        section_spec=section_spec,
        matched_chunks=matched_chunks,
        alignment=alignment,
    )


def _render_structure(
    document: Document,
    output_data: dict[str, Any],
    chunks: list[dict[str, Any]],
    style_config: dict[str, Any],
) -> list[dict[str, Any]]:
    structure = output_data.get("document_structure_proposal", {})
    sections = structure.get("sections", [])
    if not isinstance(sections, list):
        raise ValueError("document_structure_proposal.sections harus berupa list.")

    rendered: list[dict[str, Any]] = []
    alignment = _alignment(style_config.get("paragraph_alignment"))
    all_caps = bool(style_config.get("heading_all_caps", True))
    anchors = {
        index: _bookmark_name(index, section_spec)
        for index, section_spec in enumerate(sections)
        if isinstance(section_spec, dict)
    }

    for index, section_spec in enumerate(sections):
        if not isinstance(section_spec, dict):
            continue
        if index > 0:
            document.add_page_break()

        heading = _heading_text(section_spec, all_caps=all_caps)
        _add_center_heading(document, heading, bookmark_name=anchors[index], bookmark_id=index + 1)
        if str(section_spec.get("type", "")) == "bab":
            document.add_paragraph()
        matched_chunks = _match_chunks_for_section(section_spec, chunks)
        _render_section_body(
            document,
            section_spec,
            sections=sections,
            anchors=anchors,
            all_caps=all_caps,
            matched_chunks=matched_chunks,
            alignment=alignment,
        )
        rendered.append(
            {
                "index": index + 1,
                "type": section_spec.get("type"),
                "number": section_spec.get("number"),
                "title": section_spec.get("title"),
                "rendered_heading": heading,
                "bookmark": anchors[index],
                "matched_sources": [_source_line(chunk) for chunk in matched_chunks],
            }
        )

    return rendered


def render_docx(args: argparse.Namespace) -> dict[str, Any]:
    output_data = _load_metadata_payload(args.source_doc)
    chunks_payload = _load_json(Path(args.chunks))
    reasoning_report = _load_json(Path(args.reasoning_report))
    mapping = _accepted_results(reasoning_report)
    chunks = chunks_payload if isinstance(chunks_payload, list) else []

    document = Document()
    applied_page_settings = _apply_page_settings(document, mapping)
    style_config = _apply_styles(document, mapping)
    rendered_sections = _render_structure(document, output_data, chunks, style_config)

    docx_output = Path(args.output)
    docx_output.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(docx_output)
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_output = docx_output.with_name(f"{docx_output.stem}_{timestamp}{docx_output.suffix}")
        document.save(docx_output)

    render_report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "source_doc": args.source_doc,
        "metadata_source": f"document_metadata.payload::{args.source_doc}",
        "chunks": str(Path(args.chunks).resolve()),
        "reasoning_report": str(Path(args.reasoning_report).resolve()),
        "docx_output": str(docx_output.resolve()),
        "style_config": style_config,
        "applied_page_settings": applied_page_settings,
        "rendered_sections": rendered_sections,
        "note": "Sandbox render only. Tidak memakai atau mengubah docx_renderer.py.",
    }

    render_report_path = Path(args.render_report)
    render_report_path.parent.mkdir(parents=True, exist_ok=True)
    render_report_path.write_text(json.dumps(render_report, ensure_ascii=False, indent=2), encoding="utf-8")
    return render_report


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Testing ekstraksi: render DOCX sandbox dari reasoning report dan document_metadata.payload."
    )
    parser.add_argument(
        "--source-doc",
        required=True,
        help="Nama file PDF sumber yang dipakai sebagai selector document_metadata.",
    )
    parser.add_argument("--chunks", default=str(DEFAULT_CHUNKS), help="Path output_chunks.json.")
    parser.add_argument("--reasoning-report", default=str(DEFAULT_REASONING_REPORT), help="Path reasoning report.")
    parser.add_argument("--output", default=str(DEFAULT_DOCX_OUTPUT), help="Path DOCX hasil render sandbox.")
    parser.add_argument("--render-report", default=str(DEFAULT_RENDER_REPORT), help="Path report render JSON.")
    return parser


def main() -> None:
    report = render_docx(build_parser().parse_args())
    print("[reasoning-docx-render] selesai")
    print(json.dumps({"docx_output": report["docx_output"], "sections": len(report["rendered_sections"])}, indent=2))


if __name__ == "__main__":
    main()
