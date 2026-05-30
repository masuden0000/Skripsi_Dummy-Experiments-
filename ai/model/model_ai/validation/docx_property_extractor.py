"""Mengekstrak properti formatting dari file DOCX menggunakan python-docx. Posisi pipeline: DOCX input → docx_property_extractor → rule_validator."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import re as _re

from model_ai.validation.models import DocxProperties, HeadingCapsAnomaly, SpacingAnomaly


INDONESIAN_CONJUNCTIONS = {
    "di", "ke", "dari", "dan", "atau", "yang", "untuk", "dengan",
    "pada", "serta", "oleh", "dalam", "akan", "agar", "bagi",
    "sebagai", "tentang", "hingga", "namun", "tetapi", "maupun",
    "atas", "antara", "bahwa", "karena", "jika", "maka",
}

SKIP_SPACING_STYLES = ("TOC", "Caption", "Header", "Footer", "Footnote", "Table")

MAX_SPACING_ANOMALIES = 50


def _is_all_caps(text: str) -> bool:
    """Cek apakah teks sepenuhnya huruf kapital (abaikan angka dan tanda baca)."""
    alpha = [c for c in text if c.isalpha()]
    return bool(alpha) and all(c.isupper() for c in alpha)


def _is_indonesian_title_case(text: str) -> bool:
    """Cek apakah teks mengikuti title case versi Bahasa Indonesia.

    Aturan:
    - Kata pertama SELALU kapital meski kata sambung
    - Kata sambung/preposisi di tengah → huruf kecil
    - Semua kata lain → kapital di huruf pertama
    - Awalan angka / "BAB N" diabaikan sebelum cek
    """
    clean = _re.sub(
        r'^(BAB\s+\S+|[\d]+(?:\.[\d]+)*\.?)\s*',
        '',
        text.strip(),
        flags=_re.IGNORECASE,
    ).strip()
    words = clean.split()
    if not words:
        return True
    for i, w in enumerate(words):
        core = w.strip("()[].,:;\"'")
        if not core or not core[0].isalpha():
            continue
        if i == 0:
            if not core[0].isupper():
                return False
        elif core.lower() in INDONESIAN_CONJUNCTIONS:
            if core[0].isupper():
                return False
        else:
            if not core[0].isupper():
                return False
    return True


def _to_indonesian_title_case(text: str) -> str:
    """Konversi teks ke title case Bahasa Indonesia."""
    prefix_match = _re.match(
        r'^(BAB\s+\S+|[\d]+(?:\.[\d]+)*\.?)\s*',
        text.strip(),
        flags=_re.IGNORECASE,
    )
    prefix = prefix_match.group(0) if prefix_match else ""
    rest = text.strip()[len(prefix):].strip()
    words = rest.split()
    result = []
    for i, w in enumerate(words):
        if i == 0:
            result.append(w[0].upper() + w[1:] if w else w)
        elif w.lower() in INDONESIAN_CONJUNCTIONS:
            result.append(w.lower())
        else:
            result.append(w[0].upper() + w[1:] if w else w)
    return (prefix + " ".join(result)).strip()


PAPER_SIZE_MAP: dict[tuple[float, float], str] = {
    (21.0, 29.7): "A4",
    (21.0, 33.0): "F4",
    (14.85, 21.0): "A5",
    (29.7, 42.0): "A3",
    (21.59, 27.94): "LETTER",
}


def _get_paper_size(width_cm: float, height_cm: float) -> str:
    """Map paper dimensions to standard paper size names."""
    if width_cm < height_cm:
        key = (round(width_cm, 1), round(height_cm, 1))
        if key in PAPER_SIZE_MAP:
            return PAPER_SIZE_MAP[key]
    else:
        key = (round(height_cm, 1), round(width_cm, 1))
        if key in PAPER_SIZE_MAP:
            return PAPER_SIZE_MAP[key]
    return f"{round(width_cm, 1)}x{round(height_cm, 1)}cm"


def _get_alignment_string(alignment) -> str:
    """Convert python-docx alignment enum to string."""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    }
    return mapping.get(alignment, "LEFT")


def _get_orientation_string(orientation) -> str:
    """Convert python-docx orientation enum to string."""
    mapping = {
        WD_ORIENT.PORTRAIT: "PORTRAIT",
        WD_ORIENT.LANDSCAPE: "LANDSCAPE",
    }
    return mapping.get(orientation, "PORTRAIT")


def _get_line_spacing_rule(line_spacing) -> str:
    """Deprecated stub — gunakan _resolve_line_spacing_info."""
    return "MULTIPLE"


def _resolve_font_size(run, para_style, doc: Document) -> int | None:
    """Resolve effective font size dengan menelusuri style hierarchy."""
    if run.font.size is not None:
        return int(run.font.size.pt)
    style = para_style
    while style is not None:
        if style.font.size is not None:
            return int(style.font.size.pt)
        style = getattr(style, "base_style", None)
    try:
        normal = doc.styles["Normal"]
        if normal.font.size is not None:
            return int(normal.font.size.pt)
    except KeyError:
        pass
    return None


def _resolve_font_name(run, para_style, doc: Document) -> str | None:
    """Resolve effective font name dengan menelusuri style hierarchy."""
    if run.font.name is not None:
        return run.font.name
    style = para_style
    while style is not None:
        if style.font.name is not None:
            return style.font.name
        style = getattr(style, "base_style", None)
    try:
        normal = doc.styles["Normal"]
        if normal.font.name is not None:
            return normal.font.name
    except KeyError:
        pass
    return None


def _extract_typography(doc: Document) -> dict:
    """Extract typography properties dari actual paragraphs dan runs."""
    from collections import Counter

    body_sizes: Counter[int] = Counter()
    body_fonts: Counter[str] = Counter()
    heading_sizes: Counter[int] = Counter()
    heading_bold_votes: list[bool] = []
    heading_caps_votes: list[bool] = []

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        is_heading = para.style.name.startswith("Heading")

        for run in para.runs:
            if not run.text.strip():
                continue

            size = _resolve_font_size(run, para.style, doc)
            name = _resolve_font_name(run, para.style, doc)

            if size:
                if is_heading:
                    heading_sizes[size] += 1
                else:
                    body_sizes[size] += 1

            if name and not is_heading:
                body_fonts[name] += 1

        if is_heading:
            bold_val = para.runs[0].font.bold if para.runs else None
            if bold_val is None:
                bold_val = para.style.font.bold
            if bold_val is not None:
                heading_bold_votes.append(bool(bold_val))

            caps_val = para.runs[0].font.all_caps if para.runs else None
            if caps_val is None:
                caps_val = para.style.font.all_caps
            if caps_val is not None:
                heading_caps_votes.append(bool(caps_val))
            else:
                heading_caps_votes.append(para.text.strip() == para.text.strip().upper()
                                          and len(para.text.strip()) > 2)

    font_family = body_fonts.most_common(1)[0][0] if body_fonts else None
    font_size_body = body_sizes.most_common(1)[0][0] if body_sizes else None
    font_size_heading = heading_sizes.most_common(1)[0][0] if heading_sizes else None
    heading_bold = (sum(heading_bold_votes) > len(heading_bold_votes) / 2) if heading_bold_votes else None
    heading_all_caps = (sum(heading_caps_votes) > len(heading_caps_votes) / 2) if heading_caps_votes else None

    return {
        "font_family": font_family,
        "font_size_body_pt": font_size_body,
        "font_size_heading_pt": font_size_heading,
        "heading_bold": heading_bold,
        "heading_all_caps": heading_all_caps,
    }


def _twips_to_cm(value_str: str | None) -> float | None:
    """Convert twips string (may be float) to cm. 1 inch = 1440 twips = 2.54 cm."""
    if value_str is None:
        return None
    try:
        return float(value_str) / 1440 * 2.54
    except (ValueError, TypeError):
        return None


def _extract_page_layout(doc: Document) -> dict:
    """Extract page layout properties from first section."""
    section = doc.sections[0]

    pgMar = section._sectPr.find(qn("w:pgMar"))
    if pgMar is not None:
        margin_top    = _twips_to_cm(pgMar.get(qn("w:top")))
        margin_bottom = _twips_to_cm(pgMar.get(qn("w:bottom")))
        margin_left   = _twips_to_cm(pgMar.get(qn("w:left")))
        margin_right  = _twips_to_cm(pgMar.get(qn("w:right")))
    else:
        margin_top = margin_bottom = margin_left = margin_right = None

    pgSz = section._sectPr.find(qn("w:pgSz"))
    if pgSz is not None:
        page_width_cm  = _twips_to_cm(pgSz.get(qn("w:w")))
        page_height_cm = _twips_to_cm(pgSz.get(qn("w:h")))
    else:
        page_width_cm = page_height_cm = None

    paper_size = None
    if page_width_cm and page_height_cm:
        paper_size = _get_paper_size(page_width_cm, page_height_cm)

    orient_attr = pgSz.get(qn("w:orient")) if pgSz is not None else None
    if orient_attr == "landscape":
        orientation = "landscape"
    else:
        orientation = "portrait"

    cols_el = section._sectPr.find(qn("w:cols"))
    columns = 1
    if cols_el is not None:
        num_str = cols_el.get(qn("w:num"))
        if num_str:
            try:
                columns = int(num_str)
            except ValueError:
                columns = 1

    return {
        "margin_top_cm": round(margin_top, 2) if margin_top else None,
        "margin_bottom_cm": round(margin_bottom, 2) if margin_bottom else None,
        "margin_left_cm": round(margin_left, 2) if margin_left else None,
        "margin_right_cm": round(margin_right, 2) if margin_right else None,
        "paper_size": paper_size,
        "orientation": orientation,
        "columns": columns,
    }


def _resolve_line_spacing_info(para) -> tuple[float | None, str]:
    """Baca spasi baris dari XML paragraf.

    Returns:
        (line_spacing, line_spacing_rule) di mana:
        - Grup A (SINGLE/ONE_POINT_FIVE/DOUBLE): line_spacing = None, rule = nama grup
        - Grup B (MULTIPLE): line_spacing = desimal pengali, rule = "MULTIPLE"
        - Grup C (AT_LEAST/EXACTLY): line_spacing = nilai pt, rule = "AT_LEAST"/"EXACTLY"
    """
    from docx.oxml.ns import qn as _qn
    pPr = para._element.find(_qn("w:pPr"))
    if pPr is None:
        return None, "MULTIPLE"
    spacing_el = pPr.find(_qn("w:spacing"))
    if spacing_el is None:
        return None, "MULTIPLE"
    line_val = spacing_el.get(_qn("w:line"))
    line_rule = spacing_el.get(_qn("w:lineRule"))
    if line_val is None:
        return None, "MULTIPLE"

    line_int = int(line_val)

    if line_rule in (None, "auto"):
        multiplier = round(line_int / 240, 2)
        if abs(multiplier - 1.0) <= 0.02:
            return None, "SINGLE"
        elif abs(multiplier - 1.5) <= 0.02:
            return None, "ONE_POINT_FIVE"
        elif abs(multiplier - 2.0) <= 0.02:
            return None, "DOUBLE"
        return multiplier, "MULTIPLE"
    elif line_rule == "exact":
        return round(line_int / 20, 2), "EXACTLY"
    elif line_rule == "atLeast":
        return round(line_int / 20, 2), "AT_LEAST"
    else:
        multiplier = round(line_int / 240, 2)
        return multiplier, "MULTIPLE"


def _resolve_line_spacing(para) -> float | None:
    """Kompatibilitas mundur — gunakan _resolve_line_spacing_info."""
    value, _ = _resolve_line_spacing_info(para)
    return value


def _build_paragraph_map(doc: Document) -> list[dict]:
    """Single-pass: bangun peta seluruh paragraf beserta konteks lokasi hierarkisnya.

    Setiap entry berisi: index, style, text, heading_level, location, line_spacing.
    """
    para_map: list[dict] = []
    h1_ctx: str | None = None
    h2_ctx: str | None = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name

        level: int | None = None
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                pass

        if level is None and _re.match(r"^BAB\s+[IVXLC\d]+", text, _re.IGNORECASE):
            level = 1
        elif level is None and _re.match(r"^\d+\.\d+\s+\S", text):
            level = 2

        if level == 1:
            h1_ctx = text
            h2_ctx = None
        elif level == 2:
            h2_ctx = text

        if h1_ctx and h2_ctx:
            location = f"{h1_ctx} > {h2_ctx}"
        elif h1_ctx:
            location = h1_ctx
        else:
            location = f"Paragraf ke-{i + 1}"

        para_map.append({
            "index": i,
            "style": style,
            "text": text,
            "heading_level": level,
            "location": location,
            "line_spacing": _resolve_line_spacing(para),
        })

    return para_map


def _extract_heading_caps_anomalies(para_map: list[dict]) -> list[HeadingCapsAnomaly]:
    """Periksa setiap heading: H1 harus ALL CAPS, H2 harus title case Indonesia."""
    anomalies: list[HeadingCapsAnomaly] = []
    for p in para_map:
        lvl = p["heading_level"]
        text = p["text"]
        if lvl == 1:
            if not _is_all_caps(text):
                anomalies.append(HeadingCapsAnomaly(
                    text=text,
                    level=1,
                    issue="not_all_caps",
                    expected_form=text.upper(),
                    location=text,
                ))
        elif lvl == 2:
            if not _is_indonesian_title_case(text):
                anomalies.append(HeadingCapsAnomaly(
                    text=text,
                    level=2,
                    issue="not_title_case",
                    expected_form=_to_indonesian_title_case(text),
                    location=p["location"],
                ))
    return anomalies


def _extract_spacing_anomalies(
    para_map: list[dict],
    majority_spacing: float | None,
) -> list[SpacingAnomaly]:
    """Kumpulkan paragraf body yang line spacing-nya menyimpang dari majority_spacing."""
    if majority_spacing is None:
        return []
    anomalies: list[SpacingAnomaly] = []
    for p in para_map:
        if p["heading_level"] is not None:
            continue
        if any(p["style"].startswith(skip) for skip in SKIP_SPACING_STYLES):
            continue
        actual = p["line_spacing"]
        if actual is None:
            continue
        if abs(actual - majority_spacing) > 0.05:
            anomalies.append(SpacingAnomaly(
                location=p["location"],
                paragraph_index=p["index"],
                expected=majority_spacing,
                actual=actual,
                text_preview=p["text"][:60],
            ))
        if len(anomalies) >= MAX_SPACING_ANOMALIES:
            break
    return anomalies


def _detect_daftar_sections(para_map: list[dict]) -> dict[str, bool]:
    """Deteksi keberadaan bagian Daftar Isi / Pustaka / Tabel / Gambar."""
    flags: dict[str, bool] = {
        "has_daftar_isi": False,
        "has_daftar_pustaka": False,
        "has_daftar_tabel": False,
        "has_daftar_gambar": False,
    }
    for p in para_map:
        upper = p["text"].upper()
        if "DAFTAR ISI" in upper:
            flags["has_daftar_isi"] = True
        if "DAFTAR PUSTAKA" in upper or "DAFTAR REFERENSI" in upper:
            flags["has_daftar_pustaka"] = True
        if "DAFTAR TABEL" in upper:
            flags["has_daftar_tabel"] = True
        if "DAFTAR GAMBAR" in upper or "DAFTAR ILUSTRASI" in upper:
            flags["has_daftar_gambar"] = True
        if all(flags.values()):
            break
    return flags


def _extract_spacing(doc: Document) -> dict:
    """Extract spacing properties dari actual paragraphs (bukan hanya Normal style)."""
    from collections import Counter
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rule_votes: Counter[str] = Counter()
    value_by_rule: dict[str, Counter] = {}
    alignment_votes: Counter[str] = Counter()
    indent_votes: Counter[float] = Counter()

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if para.style.name.startswith("Heading"):
            continue

        ls_val, ls_rule = _resolve_line_spacing_info(para)
        rule_votes[ls_rule] += 1
        if ls_val is not None:
            valid = (
                (ls_rule == "MULTIPLE" and 0.8 <= ls_val <= 4.0)
                or (ls_rule in ("EXACTLY", "AT_LEAST") and 6.0 <= ls_val <= 144.0)
            )
            if valid:
                value_by_rule.setdefault(ls_rule, Counter())[ls_val] += 1

        align = para.alignment
        if align is None:
            align = para.style.paragraph_format.alignment
        if align is not None:
            alignment_votes[_get_alignment_string(align)] += 1

        fi = para.paragraph_format.first_line_indent
        if fi is not None:
            try:
                indent_votes[round(fi.cm, 2)] += 1
            except AttributeError:
                pass

    dominant_rule = rule_votes.most_common(1)[0][0] if rule_votes else "MULTIPLE"

    _GRUP_A = {"SINGLE", "ONE_POINT_FIVE", "DOUBLE"}
    if dominant_rule in _GRUP_A:
        line_spacing = None
    else:
        counts = value_by_rule.get(dominant_rule)
        line_spacing = counts.most_common(1)[0][0] if counts else None

    paragraph_alignment = alignment_votes.most_common(1)[0][0] if alignment_votes else "JUSTIFY"
    first_line_indent = indent_votes.most_common(1)[0][0] if indent_votes else None

    return {
        "line_spacing": line_spacing,
        "line_spacing_rule": dominant_rule,
        "paragraph_alignment": paragraph_alignment,
        "first_line_indent_cm": first_line_indent,
    }


def _extract_document_structure(doc: Document) -> dict:
    """Extract document structure dari seluruh paragraf dokumen."""
    heading_count = 0
    section_count = len(doc.sections)
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))

    return {
        "heading_count": heading_count,
        "section_count": section_count,
    }


def _extract_figures_tables(doc: Document) -> dict:
    """Extract figures and tables information."""
    table_count = len(doc.tables)

    figure_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            figure_count += 1

    figure_format = None
    table_format = None
    for para in doc.paragraphs:
        txt_p = para.text.strip()
        if not txt_p:
            continue
        style_name = para.style.name.lower()
        if "caption" in style_name or "keterangan" in style_name:
            import re as _re
            if _re.match(r"^(gambar|figure)\s+\d+[\-]\d+", txt_p, _re.IGNORECASE):
                figure_format = "Gambar {bab}.{n}"
            elif _re.match(r"^(gambar|figure)\s+\d+", txt_p, _re.IGNORECASE):
                figure_format = "Gambar {n}"
            if _re.match(r"^(tabel|table)\s+\d+[\-]\d+", txt_p, _re.IGNORECASE):
                table_format = "Tabel {bab}.{n}"
            elif _re.match(r"^(tabel|table)\s+\d+", txt_p, _re.IGNORECASE):
                table_format = "Tabel {n}"

    return {
        "table_count": table_count,
        "figure_count": figure_count,
        "figure_format": figure_format,
        "table_format": table_format,
    }


def _detect_page_number_info(doc: Document) -> dict:
    """Deteksi format dan lokasi nomor halaman dari header/footer."""
    import re as _re

    preliminary_format = None
    preliminary_location = None
    preliminary_alignment = None
    content_format = None
    content_location = None
    content_alignment = None

    def _parse_fld_instrText(xml_el) -> str | None:
        """Ambil teks instruksi field dari elemen XML."""
        for child in xml_el.iter():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "instrText" and child.text:
                return child.text.strip()
        return None

    def _detect_numfmt(xml_el) -> str | None:
        """Deteksi format angka: roman (i/I) atau arabic (1)."""
        instr = _parse_fld_instrText(xml_el)
        if instr and "PAGE" in instr:
            if "\\* Roman" in instr or "\\* roman" in instr:
                return "roman"
            if "\\* Arabic" in instr or "\\* arabic" in instr:
                return "arabic"
            return "arabic"
        return None

    def _detect_alignment(para) -> str | None:
        align = para.alignment
        if align is None and para.style:
            align = para.style.paragraph_format.alignment
        return _get_alignment_string(align) if align is not None else None

    for section in doc.sections:
        for hf_attr, location_label in (
            ("footer", "bottom"),
            ("header", "top"),
        ):
            hf = getattr(section, hf_attr, None)
            if hf is None:
                continue
            for para in hf.paragraphs:
                txt = para.text.strip()
                num_fmt = _detect_numfmt(para._element)
                if num_fmt is None:
                    if _re.search(r"\b(i|ii|iii|iv|v|vi|vii|viii|ix|x)\b", txt, _re.IGNORECASE):
                        num_fmt = "roman"
                    elif _re.search(r"\b\d+\b", txt):
                        num_fmt = "arabic"
                if num_fmt:
                    align = _detect_alignment(para)
                    if num_fmt == "roman" and preliminary_format is None:
                        preliminary_format = num_fmt
                        preliminary_location = location_label
                        preliminary_alignment = align
                    elif num_fmt == "arabic" and content_format is None:
                        content_format = num_fmt
                        content_location = location_label
                        content_alignment = align

    return {
        "preliminary_page_format": preliminary_format,
        "preliminary_page_location": preliminary_location,
        "preliminary_page_alignment": preliminary_alignment,
        "content_page_format": content_format,
        "content_page_location": content_location,
        "content_page_alignment": content_alignment,
    }


def _extract_numbering(doc: Document) -> dict:
    """Extract numbering properties dari heading patterns dan header/footer."""
    import re as _re

    chapter_format = None
    sub_chapter_format = None
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        style_name = para.style.name
        if style_name == "Heading 1" or (style_name.startswith("Heading") and "1" in style_name):
            if _re.match(r"^BAB\s+[IVXLC\d]+", txt, _re.IGNORECASE):
                chapter_format = "BAB {n}"
            elif _re.match(r"^\d+\s+\w", txt):
                chapter_format = "{n}"
        if style_name == "Heading 2" or (style_name.startswith("Heading") and "2" in style_name):
            if _re.match(r"^\d+\.\d+", txt):
                sub_chapter_format = "{n}.{m}"

    figure_format = None
    table_format = None
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        style_name = para.style.name.lower()
        if "caption" in style_name or "keterangan" in style_name:
            if _re.match(r"^(gambar|figure)\s+\d+[\.\-]\d+", txt, _re.IGNORECASE):
                figure_format = "Gambar {bab}.{n}"
            elif _re.match(r"^(gambar|figure)\s+\d+", txt, _re.IGNORECASE):
                figure_format = "Gambar {n}"
            if _re.match(r"^(tabel|table)\s+\d+[\.\-]\d+", txt, _re.IGNORECASE):
                table_format = "Tabel {bab}.{n}"
            elif _re.match(r"^(tabel|table)\s+\d+", txt, _re.IGNORECASE):
                table_format = "Tabel {n}"

    page_info = _detect_page_number_info(doc)

    return {
        "chapter_format": chapter_format,
        "sub_chapter_format": sub_chapter_format,
        **page_info,
    }


def extract_docx_properties(docx_path: str | Path) -> DocxProperties:
    """Extract all formatting properties from a DOCX file.

    Args:
        docx_path: Path to the DOCX file to extract properties from.

    Returns:
        DocxProperties object containing all extracted properties.

    Raises:
        FileNotFoundError: If the DOCX file does not exist.
        Exception: If there is an error reading the DOCX file.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {e}")

    para_map = _build_paragraph_map(doc)

    typography = _extract_typography(doc)
    page_layout = _extract_page_layout(doc)
    spacing = _extract_spacing(doc)
    doc_structure = _extract_document_structure(doc)
    figures_tables = _extract_figures_tables(doc)
    numbering = _extract_numbering(doc)

    heading_caps_anomalies = _extract_heading_caps_anomalies(para_map)
    majority_spacing = spacing["line_spacing"]
    spacing_anomalies = _extract_spacing_anomalies(para_map, majority_spacing)
    daftar_flags = _detect_daftar_sections(para_map)

    props = DocxProperties(
        font_family=typography["font_family"],
        font_size_body_pt=typography["font_size_body_pt"],
        font_size_heading_pt=typography["font_size_heading_pt"],
        heading_bold=typography["heading_bold"],
        heading_all_caps=typography["heading_all_caps"],
        margin_top_cm=page_layout["margin_top_cm"],
        margin_bottom_cm=page_layout["margin_bottom_cm"],
        margin_left_cm=page_layout["margin_left_cm"],
        margin_right_cm=page_layout["margin_right_cm"],
        paper_size=page_layout["paper_size"],
        orientation=page_layout["orientation"],
        columns=page_layout["columns"],
        line_spacing=spacing["line_spacing"],
        line_spacing_rule=spacing["line_spacing_rule"],
        paragraph_alignment=spacing["paragraph_alignment"],
        first_line_indent_cm=spacing["first_line_indent_cm"],
        heading_count=doc_structure["heading_count"],
        section_count=doc_structure["section_count"],
        has_daftar_isi=daftar_flags["has_daftar_isi"],
        has_daftar_pustaka=daftar_flags["has_daftar_pustaka"],
        has_daftar_tabel=daftar_flags["has_daftar_tabel"],
        has_daftar_gambar=daftar_flags["has_daftar_gambar"],
        chapter_format=numbering["chapter_format"],
        sub_chapter_format=numbering["sub_chapter_format"],
        preliminary_page_format=numbering["preliminary_page_format"],
        preliminary_page_location=numbering["preliminary_page_location"],
        preliminary_page_alignment=numbering["preliminary_page_alignment"],
        content_page_format=numbering["content_page_format"],
        content_page_location=numbering["content_page_location"],
        content_page_alignment=numbering["content_page_alignment"],
        figure_format=figures_tables["figure_format"],
        table_format=figures_tables["table_format"],
        table_count=figures_tables["table_count"],
        figure_count=figures_tables["figure_count"],
        heading_caps_anomalies=heading_caps_anomalies,
        spacing_anomalies=spacing_anomalies,
    )

    return props


def extract_docx_properties_dict(docx_path: str | Path) -> dict:
    """Extract all formatting properties from a DOCX file as dict.

    Args:
        docx_path: Path to the DOCX file to extract properties from.

    Returns:
        Dictionary containing all extracted properties.
    """
    props = extract_docx_properties(docx_path)
    return props.to_dict()