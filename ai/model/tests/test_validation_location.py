"""Test deteksi page dan bab di _get_para_details()."""
import os
import tempfile

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Impor fungsi yang akan diuji
from model_ai.validation.validocx.debug_report import _get_para_details


def _make_test_docx_with_pages() -> str:
    """Buat DOCX sementara: 2 halaman, 2 BAB, masing-masing 1 paragraf isi."""
    doc = DocxDocument()

    # Halaman 1 — BAB 1
    doc.add_heading("BAB 1 PENDAHULUAN", level=1)
    doc.add_paragraph("Paragraf isi BAB 1.")

    # Explicit page break → halaman 2
    para_break = doc.add_paragraph()
    run_break = para_break.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_break._r.append(br)

    # Halaman 2 — BAB 2
    doc.add_heading("BAB 2 TINJAUAN PUSTAKA", level=1)
    doc.add_paragraph("Paragraf isi BAB 2.")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def test_page_and_bab_are_detected():
    path = _make_test_docx_with_pages()
    try:
        result = _get_para_details(path)

        # Cari paragraf isi BAB 1
        bab1_para = next(
            v for v in result.values() if v.get("text") == "Paragraf isi BAB 1."
        )
        assert bab1_para["page"] == 1, "Paragraf BAB 1 harus di halaman 1"
        assert bab1_para["bab"] == "BAB 1 PENDAHULUAN"

        # Cari paragraf isi BAB 2
        bab2_para = next(
            v for v in result.values() if v.get("text") == "Paragraf isi BAB 2."
        )
        assert bab2_para["page"] == 2, "Paragraf BAB 2 harus di halaman 2"
        assert bab2_para["bab"] == "BAB 2 TINJAUAN PUSTAKA"
    finally:
        os.unlink(path)


def test_bab_is_none_before_first_heading():
    doc = DocxDocument()
    doc.add_paragraph("Paragraf sebelum BAB manapun.")
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    try:
        result = _get_para_details(tmp.name)
        first = next(
            v for v in result.values()
            if v.get("text") == "Paragraf sebelum BAB manapun."
        )
        assert first["bab"] is None
        assert first["page"] == 1
    finally:
        os.unlink(tmp.name)


from model_ai.validation.models import ValidationIssue


def test_validation_issue_accepts_occurrences():
    issue = ValidationIssue(
        category="typography",
        field="font_per_paragraph",
        severity="error",
        message="Font mismatch",
        occurrences=[
            {
                "page": 3,
                "bab": "BAB 1 PENDAHULUAN",
                "para_idx": 4,
                "style": "Normal",
                "text": "Latar belakang...",
                "actual": "11.0pt, Times New Roman",
                "expected": "12pt, Times New Roman",
            }
        ],
    )
    assert issue.occurrences is not None
    assert len(issue.occurrences) == 1
    assert issue.occurrences[0]["page"] == 3
    assert issue.occurrences[0]["bab"] == "BAB 1 PENDAHULUAN"


def test_validation_issue_occurrences_defaults_to_none():
    issue = ValidationIssue(
        category="typography",
        field="font_per_paragraph",
        severity="error",
        message="Font mismatch",
    )
    assert issue.occurrences is None
