"""Test bahwa iter_paragraphs() membaca paragraf di dalam w:sdt (mis. TOC Word)."""
from pathlib import Path
from docx import Document
from model_ai.validation.validocx.wrapper import DocumentWrapper

_DOCX = Path(__file__).parent / "file_target.docx"


def test_iter_paragraphs_includes_sdt_paragraphs():
    """Jumlah paragraf dari iter_paragraphs() harus lebih banyak dari doc.paragraphs
    karena dokumen punya TOC di dalam w:sdt."""
    doc = Document(str(_DOCX))
    wrapper = DocumentWrapper(doc)

    standard_count = len(list(doc.paragraphs))
    iter_count = sum(1 for _ in wrapper.iter_paragraphs())

    assert iter_count > standard_count, (
        f"iter_paragraphs ({iter_count}) harus lebih banyak dari "
        f"doc.paragraphs ({standard_count}) karena ada w:sdt di dokumen"
    )


def test_iter_paragraphs_yields_toc_styles():
    """iter_paragraphs() harus menghasilkan setidaknya satu paragraf dengan style mengandung 'TOC'."""
    doc = Document(str(_DOCX))
    wrapper = DocumentWrapper(doc)

    toc_styles_found = [
        p.style.name for p in wrapper.iter_paragraphs()
        if "TOC" in p.style.name or "toc" in p.style.name.lower()
    ]

    assert len(toc_styles_found) > 0, (
        "iter_paragraphs() tidak menghasilkan paragraf dengan style TOC — "
        "kemungkinan w:sdt tidak ditelusuri"
    )
